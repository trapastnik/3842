# -*- coding: utf-8 -*-
"""Каркас генерации .docx: титул, стили, блоки."""
import os

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "Times New Roman"
ORG = "ЛЕНИНСКИЙ МЕМОРИАЛ"
CITY = "г. Санкт-Петербург"
YEAR = "2026 г."

# Титульная страница берётся из образца с логотипом музея (файл рядом с генератором):
# сохраняются логотип, шапка, город и год; подставляются только три строки —
# обозначение МТК, название документа и подзаголовок.
TEMPLATE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "титул-образец.docx")
TITLE_SLOTS = (3, 4, 5)   # номера абзацев титула: МТК · тип документа · подзаголовок


def _set_font(run, size=12, bold=False, italic=False, color=None):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(a), FONT)


def _grid_borders(table):
    """Сетка таблицы напрямую через XML — не зависит от набора стилей в шаблоне."""
    tbl_pr = table._tbl.tblPr
    for old in tbl_pr.findall(qn("w:tblBorders")):
        tbl_pr.remove(old)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        borders.append(el)
    tbl_pr.append(borders)


class Doc:
    def __init__(self, mtk_line, doc_type, subtitle=None):
        self.d = Document(TEMPLATE)
        self._keep_title_page()
        self._page_setup()
        self._base_style()
        self.h1_n = 0
        self.h2_n = 0
        self.fig_n = 0
        self.tab_n = 0
        self._title_page(mtk_line, doc_type, subtitle)

    # ---------- инфраструктура ----------
    def _page_setup(self):
        s = self.d.sections[0]
        s.left_margin = Cm(2.5)
        s.right_margin = Cm(1.5)
        s.top_margin = Cm(2.0)
        s.bottom_margin = Cm(2.0)

    def _base_style(self):
        st = self.d.styles["Normal"]
        st.font.name = FONT
        st.font.size = Pt(12)
        st.paragraph_format.line_spacing = 1.15
        st.paragraph_format.space_after = Pt(6)
        rpr = st.element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rfonts.set(qn(a), FONT)

    def _par(self, text="", size=12, bold=False, italic=False, align=None,
             space_before=0, space_after=6, color=None, indent=None):
        p = self.d.add_paragraph()
        if align is not None:
            p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        if indent is not None:
            p.paragraph_format.left_indent = Cm(indent)
        if text:
            _set_font(p.add_run(text), size, bold, italic, color)
        return p

    def _keep_title_page(self):
        """Оставить в шаблоне только титульную страницу (до первого разрыва включительно)."""
        body = self.d.element.body
        kids = [ch for ch in body.iterchildren()]
        end = None
        for i, ch in enumerate(kids):
            if ch.tag != qn("w:p"):
                continue
            for br in ch.findall(".//" + qn("w:br")):
                if br.get(qn("w:type")) == "page":
                    end = i
                    break
            if end is not None:
                break
        if end is None:
            raise RuntimeError("в образце титула не найден разрыв страницы")
        for ch in kids[end + 1:]:
            if ch.tag == qn("w:sectPr"):
                continue
            body.remove(ch)
        self._title_pars = [ch for ch in kids[:end + 1] if ch.tag == qn("w:p")]

    def _title_page(self, mtk_line, doc_type, subtitle):
        """Подставить в титул образца обозначение МТК, название документа и подзаголовок."""
        values = (mtk_line, doc_type.upper(), subtitle or "")
        for slot, value in zip(TITLE_SLOTS, values):
            p = self._title_pars[slot]
            ts = p.findall(".//" + qn("w:t"))
            if not ts:
                raise RuntimeError(f"в абзаце титула {slot} нет текста для замены")
            ts[0].text = value
            ts[0].set(qn("xml:space"), "preserve")
            for extra in ts[1:]:
                extra.text = ""

    def page_break(self):
        self.d.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # ---------- блоки ----------
    def annotation(self, paragraphs):
        self._par("АННОТАЦИЯ", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
        for t in paragraphs:
            self._par(t, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        self.page_break()

    def toc(self):
        self._par("СОДЕРЖАНИЕ", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
        p = self.d.add_paragraph()
        r = p.add_run()
        _set_font(r, 12)
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), r'TOC \o "1-2" \h \z \u')
        inner = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = "Оглавление формируется в Word: «Ссылки» → «Обновить таблицу»."
        inner.append(t)
        fld.append(inner)
        p._p.append(fld)
        self.page_break()

    def h1(self, text):
        self.h1_n += 1
        self.h2_n = 0
        p = self.d.add_paragraph(style="Heading 1")
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(8)
        _set_font(p.add_run(f"{self.h1_n}. {text}"), 14, bold=True, color=(0, 0, 0))
        return p

    def h2(self, text):
        self.h2_n += 1
        p = self.d.add_paragraph(style="Heading 2")
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(6)
        _set_font(p.add_run(f"{self.h1_n}.{self.h2_n}. {text}"), 12.5, bold=True, color=(0, 0, 0))
        return p

    def p(self, text):
        self._par(text, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    def ul(self, items, dash="—"):
        for it in items:
            self._par(f"{dash} {it}", align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                      indent=0.5, space_after=3)

    def ol(self, items):
        for i, it in enumerate(items, 1):
            self._par(f"{i}) {it}", align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                      indent=0.5, space_after=3)

    def table(self, header, rows, caption=None, widths=None, size=10.5):
        if caption:
            self.tab_n += 1
            self._par(f"Таблица {self.tab_n} — {caption}", size=11, italic=True,
                      align=WD_ALIGN_PARAGRAPH.RIGHT, space_before=8, space_after=3)
        t = self.d.add_table(rows=1, cols=len(header))
        _grid_borders(t)
        hdr = t.rows[0].cells
        for i, h in enumerate(header):
            hdr[i].text = ""
            pp = hdr[i].paragraphs[0]
            pp.paragraph_format.space_after = Pt(2)
            _set_font(pp.add_run(h), size, bold=True)
        for row in rows:
            cells = t.add_row().cells
            for i, v in enumerate(row):
                cells[i].text = ""
                pp = cells[i].paragraphs[0]
                pp.paragraph_format.space_after = Pt(2)
                _set_font(pp.add_run(str(v)), size)
        if widths:
            for r in t.rows:
                for i, w in enumerate(widths):
                    r.cells[i].width = Cm(w)
        self._par(space_after=4)
        return t

    def fig(self, caption, capture):
        """Место под иллюстрацию + указание, что снять."""
        self.fig_n += 1
        self._par(f"[ МЕСТО ПОД ИЛЛЮСТРАЦИЮ. Снимок экрана: {capture} ]",
                  size=11, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER,
                  color=(0x80, 0x80, 0x80), space_before=8, space_after=3)
        self._par(f"Рисунок {self.fig_n}. {caption}", size=11, italic=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)

    def note(self, text):
        self._par(text, size=11, italic=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  indent=0.5, space_after=8)

    def save(self, path):
        self.d.save(path)
